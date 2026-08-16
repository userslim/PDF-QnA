"""文档处理模块 - 支持 PDF、Word 和 OCR 识别"""
import os
import uuid
from pathlib import Path
from typing import List, Dict
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import io


class DocumentChunk:
    """文档块，包含文本内容和来源信息"""
    def __init__(self, text: str, source: str, page: int = None, chunk_id: str = None):
        self.text = text
        self.source = source
        self.page = page
        self.chunk_id = chunk_id or str(uuid.uuid4())


class DocumentProcessor:
    """文档处理器，支持多种文件格式"""
    
    def __init__(self):
        self.ocr_engine = None
        # OCR 不在启动时加载，改为首次使用时加载
        self._ocr_loaded = False
    
    def _init_ocr(self):
        """延迟加载 OCR（避免启动慢）"""
        if self._ocr_loaded:
            return
        try:
            from paddleocr import PaddleOCR
            # 使用最简参数调用，确保版本兼容
            self.ocr_engine = PaddleOCR(lang='en')
        except Exception:
            try:
                self.ocr_engine = PaddleOCR(use_angle_cls=True, lang='en')
            except Exception as e:
                print(f"OCR 初始化失败: {e}")
                self.ocr_engine = None
        self._ocr_loaded = True
    
    def process_file(self, file_path: str) -> List[DocumentChunk]:
        """根据文件类型处理文档"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.process_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.process_word(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    
    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """处理 PDF 文件，包括 OCR 识别图片中的文字"""
        chunks = []
        filename = os.path.basename(file_path)
        
        # 方法 1: 使用 pdfplumber 提取文本
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        # 分块（每页或每 1000 字符）
                        chunk_size = 1000
                        for i in range(0, len(text), chunk_size):
                            chunk_text = text[i:i+chunk_size]
                            chunks.append(DocumentChunk(
                                text=chunk_text,
                                source=filename,
                                page=page_num
                            ))
        except Exception as e:
            print(f"pdfplumber 提取失败: {e}")
        
        # 方法 2: 使用 PyMuPDF + OCR 提取图片中的文字
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 提取页面图片
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # OCR 识别
                    if self.ocr_engine:
                        try:
                            image = Image.open(io.BytesIO(image_bytes))
                            # 转换为 numpy 数组（PaddleOCR 新版本需要）
                            import numpy as np
                            img_array = np.array(image)
                            
                            # 兼容新旧版本 API
                            try:
                                # 新版本 API
                                ocr_result = self.ocr_engine.predict(img_array)
                            except AttributeError:
                                # 旧版本 API
                                ocr_result = self.ocr_engine.ocr(img_array, cls=True)
                            
                            # 解析结果（新旧版本格式不同）
                            if ocr_result:
                                ocr_text = ""
                                if isinstance(ocr_result, list) and len(ocr_result) > 0:
                                    # 旧版本格式: [[[box, (text, conf)], ...]]
                                    result_data = ocr_result[0] if isinstance(ocr_result[0], list) else ocr_result
                                    if result_data and len(result_data) > 0:
                                        ocr_text = "\n".join([
                                            line[1][0] if isinstance(line[1], tuple) else str(line[1])
                                            for line in result_data
                                        ])
                                elif isinstance(ocr_result, dict):
                                    # 新版本格式
                                    if 'rec_text' in ocr_result:
                                        ocr_text = "\n".join(ocr_result['rec_text'])
                                
                                if ocr_text.strip():
                                    chunks.append(DocumentChunk(
                                        text=f"[OCR 图片文字] {ocr_text}",
                                        source=filename,
                                        page=page_num + 1
                                    ))
                        except Exception as e:
                            print(f"OCR 识别失败 (页 {page_num + 1}, 图 {img_index}): {e}")
            
            doc.close()
        except Exception as e:
            print(f"PyMuPDF 处理失败: {e}")
        
        return chunks
    
    def process_word(self, file_path: str) -> List[DocumentChunk]:
        """处理 Word 文档"""
        chunks = []
        filename = os.path.basename(file_path)
        
        try:
            doc = Document(file_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    chunks.append(DocumentChunk(
                        text=text,
                        source=filename,
                        page=para_num + 1
                    ))
            
            # 处理表格
            for table_num, table in enumerate(doc.tables):
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        chunks.append(DocumentChunk(
                            text=row_text,
                            source=filename,
                            page=f"表格 {table_num + 1}"
                        ))
        except Exception as e:
            print(f"Word 处理失败: {e}")
        
        return chunks
    
    def process_multiple_files(self, file_paths: List[str]) -> List[DocumentChunk]:
        """批量处理多个文件"""
        all_chunks = []
        for file_path in file_paths:
            try:
                chunks = self.process_file(file_path)
                all_chunks.extend(chunks)
                print(f"已处理: {os.path.basename(file_path)} - {len(chunks)} 个块")
            except Exception as e:
                print(f"处理失败 {file_path}: {e}")
        
        return all_chunks


if __name__ == "__main__":
    # 测试
    processor = DocumentProcessor()
    print("文档处理器初始化成功")